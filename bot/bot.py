import telebot
from db.db import async_session
from telebot import types
from config import BOT_TOKEN
from auth.auth import authenticate_user, register_user
from docx import Document

global ACCESS, audio_file
global protocol_FUNC1, protocol_FUNC2, protocol_FUNC3
global unofficial_FUNC1, unofficial_FUNC2, unofficial_FUNC3
global FUNC1, FUNC2, FUNC3
bot = telebot.TeleBot(BOT_TOKEN)

bot.delete_webhook()

commands = [
        types.BotCommand(command="/start", description="Начать работу"),
        types.BotCommand(command="/help", description="Информационное сообщение"),
    ]
bot.set_my_commands(commands)

@bot.message_handler(commands=['start'])
def send_welcome(message):

    markup1 = telebot.types.InlineKeyboardMarkup(row_width=2)
    markup_item1 = telebot.types.InlineKeyboardButton('Войти', callback_data='Sign in')
    markup_item2 = telebot.types.InlineKeyboardButton('Зарегистрироваться', callback_data='Log in')
    markup1.add(markup_item1, markup_item2)

    bot.send_message(message.from_user.id, "Привет! Я бот секретарь с искусственным интеллектом", reply_markup=markup1)



@bot.message_handler(commands=['help'])
def send_welcome(message):
    text = 'Здравствуйте!\nДанный бот предназначен для  для создания протоколов совещаний и генерации их расшифровки.\n\n*Расшифровка совещания* позволяет получить в текстовом формате запись совещаний.\n\n*Официальный протокол* включает в себя: дату, время, перечень присутствующих, деление на основные блоки, ключевые предложения, контекст обсуждения каждого предложения, время в аудиозаписи, протокол поручений (кому поручено, контекст поручения, срок выполнения).\n\n*Неофициальный протокол* включает в себя: дату, время, длительность, участников (опционально), повестку дня (опционально), деление на основные блоки, контекст обсуждения каждого предложения (опционально),время в аудиозаписи.'
    bot.send_message(message.from_user.id, text, parse_mode='Markdown')


def markup_change_format_file_and_password(Format_docx:bool, Format_pdf:bool, Password:bool) -> types.InlineKeyboardMarkup:
    global FUNC1, FUNC2, FUNC3        # необходимые глобальные переменные(чтобы их видел callback_inline()) для проверки состояния чекбоксов
    FUNC1, FUNC2, FUNC3 = Format_docx, Format_pdf, Password

    markup = telebot.types.InlineKeyboardMarkup(row_width=1)       # создание пустого markup

    text_func1 = 'Формат docx: ✅' if Format_docx else 'Формат docx: ⬜'
    text_func2 = 'Формат pdf: ✅' if Format_pdf else 'Формат pdf: ⬜'            # проверка, какой чекбокс включен, а какой нет
    text_func3 = 'Установить пароль: ✅' if Password else 'Установить пароль: ⬜'
    text_func4 = 'Подтвердить'

    markup_item1 = telebot.types.InlineKeyboardButton(text_func1, callback_data='Format docx')
    markup_item2 = telebot.types.InlineKeyboardButton(text_func2, callback_data='Format pdf') # создание верного чекбокс(вкл/выкл)
    markup_item3 = telebot.types.InlineKeyboardButton(text_func3, callback_data='Password')
    markup_item4 = telebot.types.InlineKeyboardButton(text_func4, callback_data='Confirm format file')

    markup.add(markup_item1, markup_item2, markup_item3, markup_item4)

    return markup


def markup_redakt_unofficial_protocol(Meeting_participants:bool, Agendas:bool, Context_of_discussion:bool) -> types.InlineKeyboardMarkup:
    global unofficial_FUNC1, unofficial_FUNC2, unofficial_FUNC3        # необходимые глобальные переменные(чтобы их видел callback_inline()) для проверки состояния чекбоксов
    unofficial_FUNC1, unofficial_FUNC2, unofficial_FUNC3 = Meeting_participants, Agendas, Context_of_discussion

    markup = telebot.types.InlineKeyboardMarkup(row_width=1)       # создание пустого markup

    text_func1 = 'Участники совещания: ✅' if Meeting_participants else 'Участники совещания: ⬜'
    text_func2 = 'Повестка дня: ✅' if Agendas else 'Повестка дня: ⬜'            # проверка, какой чекбокс включен, а какой нет
    text_func3 = 'Контекст обсуждения каждого предложения: ✅' if Context_of_discussion else 'Контекст обсуждения каждого предложения: ⬜'
    text_func4 = 'Подтвердить'

    markup_item1 = telebot.types.InlineKeyboardButton(text_func1, callback_data='Meeting participants')
    markup_item2 = telebot.types.InlineKeyboardButton(text_func2, callback_data='Agendas') # создание верного чекбокс(вкл/выкл)
    markup_item3 = telebot.types.InlineKeyboardButton(text_func3, callback_data='Context of discussion')
    markup_item4 = telebot.types.InlineKeyboardButton(text_func4, callback_data='Confirm unofficial protocol')

    markup.add(markup_item1, markup_item2, markup_item3, markup_item4)

    return markup


def chek_markup_protocol(Unofficial_protocol:bool, Official_protocol:bool, Transcript_of_the_meeting:bool) -> types.InlineKeyboardMarkup:
    global protocol_FUNC1, protocol_FUNC2, protocol_FUNC3        # необходимые глобальные переменные(чтобы их видел callback_inline()) для проверки состояния чекбоксов
    protocol_FUNC1, protocol_FUNC2, protocol_FUNC3 = Unofficial_protocol, Official_protocol, Transcript_of_the_meeting

    markup = telebot.types.InlineKeyboardMarkup(row_width=1)       # создание пустого markup

    text_func1 = 'Неофициальный протокол: ✅' if Unofficial_protocol else 'Неофициальный протокол: ⬜'
    text_func2 = 'Официальный протокол: ✅' if Official_protocol else 'Официальный протокол: ⬜'            # проверка, какой чекбокс включен, а какой нет
    text_func3 = 'Расшифровка совещания: ✅' if Transcript_of_the_meeting else 'Расшифровка совещания: ⬜'
    text_func4 = 'Подтвердить'

    markup_item1 = telebot.types.InlineKeyboardButton(text_func1, callback_data='Unofficial protocol')
    markup_item2 = telebot.types.InlineKeyboardButton(text_func2, callback_data='Official protocol') # создание верного чекбокс(вкл/выкл)
    markup_item3 = telebot.types.InlineKeyboardButton(text_func3, callback_data='Transcript of the meeting')
    markup_item4 = telebot.types.InlineKeyboardButton(text_func4, callback_data='Confirm protocol')

    markup.add(markup_item1, markup_item2, markup_item3, markup_item4)

    return markup




def create_check_boxes(message, markup, text) -> None:
    global markup_protocol
    markup_protocol = False
    bot.send_message(message.chat.id,
                     text=text,
                     reply_markup=markup)


@bot.callback_query_handler(func=lambda call: True)
def callback_inline(call):
    if call.message:
        if call.data == 'Unofficial protocol' or call.data == 'Official protocol' or call.data == 'Transcript of the meeting':
            if call.data == 'Unofficial protocol':
                markup = chek_markup_protocol(not protocol_FUNC1, protocol_FUNC2, protocol_FUNC3)   # изменение состояния чекбокса (№ 1)
            elif call.data == 'Official protocol':
                markup = chek_markup_protocol(protocol_FUNC1, not protocol_FUNC2, protocol_FUNC3)   # изменение состояния чекбокса (№ 2)
            elif call.data == 'Transcript of the meeting':
                markup = chek_markup_protocol(protocol_FUNC1, protocol_FUNC2, not protocol_FUNC3)   # изменение состояния чекбокса (№ 3)
    
            bot.edit_message_text(chat_id=call.message.chat.id,
                                  message_id=call.message.message_id,
                                  text="Ваша аудиозапись обработана, в каком виде вам нужен протокол:",
                                  reply_markup=markup)
        if call.data == 'Meeting participants' or call.data == 'Agendas' or call.data == 'Context of discussion':
            if call.data == 'Meeting participants':
                markup = markup_redakt_unofficial_protocol(not unofficial_FUNC1, unofficial_FUNC2, unofficial_FUNC3)  # изменение состояния чекбокса (№ 1)
            elif call.data == 'Agendas':
                markup = markup_redakt_unofficial_protocol(unofficial_FUNC1, not unofficial_FUNC2, unofficial_FUNC3)  # изменение состояния чекбокса (№ 2)
            elif call.data == 'Context of discussion':
                markup = markup_redakt_unofficial_protocol(unofficial_FUNC1, unofficial_FUNC2, not unofficial_FUNC3)  # изменение состояния чекбокса (№ 3)

            bot.edit_message_text(chat_id=call.message.chat.id,
                                  message_id=call.message.message_id,
                                  text="Выберите необязательные поля для неофициального протокола:",
                                  reply_markup=markup)
        if call.data == 'Format docx' or call.data == 'Format pdf' or call.data == 'Password':
            if call.data == 'Format docx':
                markup = markup_change_format_file_and_password(not FUNC1, FUNC2, FUNC3)  # изменение состояния чекбокса (№ 1)
            elif call.data == 'Format pdf':
                markup = markup_change_format_file_and_password(FUNC1, not FUNC2, FUNC3)  # изменение состояния чекбокса (№ 2)
            elif call.data == 'Password':
                markup = markup_change_format_file_and_password(FUNC1, FUNC2, not FUNC3)  # изменение состояния чекбокса (№ 3)

            bot.edit_message_text(chat_id=call.message.chat.id,
                                  message_id=call.message.message_id,
                                  text="Выберите нужный(ые) формат(ы) и отметьте нужен ли пароль на документ:",
                                  reply_markup=markup)
        if call.data == 'Sign in':
            mes = bot.send_message(call.message.chat.id, 'Введите:\nЛогин\nПароль')
            bot.register_next_step_handler(mes, sign_up)

        if call.data == 'Log in':
            mes = bot.send_message(call.message.chat.id, 'Зарегистрируйте:\nЛогин\nПароль')
            bot.register_next_step_handler(mes, log_in)

        if call.data == 'Confirm protocol':
            if protocol_FUNC1:
                create_check_boxes(call.message, markup_redakt_unofficial_protocol(False, False, False), 'Выберите необязательные поля для неофициального протокола:')
            elif protocol_FUNC2:
                create_check_boxes(call.message, markup_change_format_file_and_password(False, False, False), 'Выберите нужный(ые) формат(ы) и отметьте нужен ли пароль на документ:')
            elif protocol_FUNC3:
                pass
        if call.data == 'Confirm unofficial protocol':
            create_check_boxes(call.message, markup_change_format_file_and_password(False, False, False),
                               'Выберите нужный(ые) формат(ы) и отметьте нужен ли пароль на документ:')
        if call.data == 'Confirm format file':
            if protocol_FUNC1 and FUNC1:
                name_file = 'Неофициальный протокол.docx'
                create_file(name_file)
                bot.send_document(call.message.chat.id, open(name_file, 'rb'), caption='Благодарим за использование, нужный(ые) вам документ(ы) прикреплен(ы) к этому сообщению')
            if protocol_FUNC2 and FUNC1:
                name_file = 'Официальный протокол.docx'
                create_file(name_file)
                bot.send_document(call.message.chat.id, open(name_file, 'rb'), caption='Благодарим за использование, нужный(ые) вам документ(ы) прикреплен(ы) к этому сообщению')



def sign_up(message):
    try:
        email = message.text.split('\n')[0]
        password = message.text.split('\n')[1]
        user_id = message.from_user.id
        print(user_id, email, password)
        authenticate_user(user_id, email, password, async_session)
        bot.send_message(message.from_user.id, 'Успешный вход')
        mes = bot.send_message(message.chat.id, 'Отправьте аудиозапись совещания:')
        bot.register_next_step_handler(mes, get_audio_file)

    except Exception:
        bot.send_message(message.from_user.id, 'Некорректный ввод')


def log_in(message):
    try:
        email = message.text.split('\n')[0]
        password = message.text.split('\n')[1]
        user_id = message.from_user.id
        print(user_id, email, password)
        register_user(user_id, email, password, async_session)
        bot.send_message(message.from_user.id, 'Успешная регистрация')


    except Exception:
        bot.send_message(message.from_user.id, 'Некорректный ввод')


def get_audio_file(message):
    global audio_file
    audio_file = message.audio
    bot.send_message(message.chat.id, 'Идет обработка...')
    create_check_boxes(message, chek_markup_protocol(False, False, False), "Ваша аудиозапись обработана, в каком виде вам нужен протокол:")


def create_file(name_file):
    doc = Document()

    doc.add_heading('Document Title', 0)

    p = doc.add_paragraph('A plain paragraph having some')
    p.add_run('bold').bold = True
    p.add_run('and some ')
    p.add_run('italic.').italic = True

    doc.add_page_break()

    doc.save(name_file)


@bot.message_handler(func=lambda message: True)
def answer(message):
    bot.send_message(message.from_user.id, 'Я вас не понимаю(')


if __name__ == '__main__':
    bot.infinity_polling(none_stop=True)



